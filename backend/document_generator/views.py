from django.shortcuts import render
from rest_framework import viewsets, status
from rest_framework.response import Response
from rest_framework.decorators import action
from django.conf import settings
from django.http import FileResponse, HttpResponse
import os
from docx import Document
import pandas as pd
import zipfile
import io
from .models import Template, DataFile, DocumentGeneration
from .serializers import TemplateSerializer, DataFileSerializer, DocumentGenerationSerializer

class TemplateViewSet(viewsets.ModelViewSet):
    queryset = Template.objects.all()
    serializer_class = TemplateSerializer

    @action(detail=True, methods=['post'])
    def extract_placeholders(self, request, pk=None):
        template = self.get_object()
        doc = Document(template.template_file.path)
        placeholders = set()
        
        for paragraph in doc.paragraphs:
            text = paragraph.text
            start = 0
            while True:
                start = text.find('{{', start)
                if start == -1:
                    break
                end = text.find('}}', start)
                if end == -1:
                    break
                placeholder = text[start+2:end].strip()
                placeholders.add(placeholder)
                start = end + 2
        
        template.placeholders = list(placeholders)
        template.save()
        return Response({'placeholders': template.placeholders})

class DataFileViewSet(viewsets.ModelViewSet):
    queryset = DataFile.objects.all()
    serializer_class = DataFileSerializer

    @action(detail=True, methods=['get'])
    def get_headers(self, request, pk=None):
        try:
            data_file = self.get_object()
            df = pd.read_excel(data_file.data_file.path)
            headers = list(df.columns)
            data_file.headers = headers
            data_file.save()
            return Response({'headers': headers})
        except Exception as e:
            return Response(
                {'detail': f'Error reading Excel file: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )

class DocumentGenerationViewSet(viewsets.ModelViewSet):
    queryset = DocumentGeneration.objects.all()
    serializer_class = DocumentGenerationSerializer

    @action(detail=True, methods=['post'])
    def generate_documents(self, request, pk=None):
        try:
            doc_gen = self.get_object()
            
            # Create a temporary directory for storing files
            temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_generated')
            os.makedirs(temp_dir, exist_ok=True)

            # Import docx2pdf for PDF conversion
            if doc_gen.output_format == 'pdf':
                from docx2pdf import convert            # Load template and data
            try:
                template_doc = Document(doc_gen.template.template_file.path)
                df = pd.read_excel(doc_gen.data_file.data_file.path)
            except Exception as e:
                return Response(
                    {'detail': f'Error loading template or data file: {str(e)}'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Handle row range if specified
            if doc_gen.start_row is not None and doc_gen.end_row is not None:
                # Python uses 0-based indexing, but we want to provide a 1-based interface to users
                start_idx = max(doc_gen.start_row - 1, 0)  # Prevent negative index
                end_idx = min(doc_gen.end_row, len(df))    # Prevent overflow
                df = df.iloc[start_idx:end_idx]
            elif doc_gen.start_row is not None:
                start_idx = max(doc_gen.start_row - 1, 0)
                df = df.iloc[start_idx:]
            elif doc_gen.end_row is not None:
                end_idx = min(doc_gen.end_row, len(df))
                df = df.iloc[:end_idx]

            # Validate mapping against headers
            for excel_column in doc_gen.mapping.values():
                if excel_column not in df.columns:
                    return Response(
                        {'detail': f'Excel column "{excel_column}" not found in data file'},
                        status=status.HTTP_400_BAD_REQUEST
                    )

            output_files = []
            for index, row in df.iterrows():
                try:
                    doc = Document(doc_gen.template.template_file.path)
                    
                    for paragraph in doc.paragraphs:
                        for key, value in doc_gen.mapping.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    '{{' + key + '}}',
                                    str(row[value])
                                )                    # Generate filename using pattern
                    filename_pattern = doc_gen.filename_pattern or 'document_{{id}}'
                    output_filename = filename_pattern.replace('{{id}}', str(index + 1))
                    
                    # Replace placeholders in filename
                    for key, value in doc_gen.mapping.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in output_filename:
                            output_filename = output_filename.replace(placeholder, str(row[value]))
                    
                    # Sanitize filename
                    output_filename = "".join(c for c in output_filename if c.isalnum() or c in ('-', '_', ' '))
                    
                    if doc_gen.output_format == 'docx':
                        output_path = os.path.join(temp_dir, f'{output_filename}.docx')
                        doc.save(output_path)
                    elif doc_gen.output_format == 'pdf':
                        # First save as docx
                        docx_path = os.path.join(temp_dir, f'{output_filename}.docx')
                        doc.save(docx_path)
                        # Convert to PDF
                        output_path = os.path.join(temp_dir, f'{output_filename}.pdf')
                        convert(docx_path, output_path)
                        # Remove temporary docx file
                        try:
                            os.remove(docx_path)
                        except:
                            pass  # Ignore if temp file can't be deleted
                    elif doc_gen.output_format == 'txt':
                        output_path = os.path.join(temp_dir, f'{output_filename}.txt')
                        with open(output_path, 'w', encoding='utf-8') as f:
                            for para in doc.paragraphs:
                                f.write(para.text + '\n')

                    output_files.append(output_path)
                except Exception as e:
                    return Response(
                        {'detail': f'Error generating document {index + 1}: {str(e)}'},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )

            # Create a zip file containing all generated documents
            zip_filename = 'generated_documents.zip'
            zip_path = os.path.join(temp_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zip_file:
                for file_path in output_files:
                    zip_file.write(file_path, os.path.basename(file_path))

            # Read the zip file into memory
            with open(zip_path, 'rb') as f:
                zip_data = f.read()

            # Clean up temporary files
            for file_path in output_files:
                try:
                    os.remove(file_path)
                except:
                    pass
            try:
                os.remove(zip_path)
            except:
                pass

            doc_gen.status = 'completed'
            doc_gen.save()

            # Return zip file as response
            response = HttpResponse(zip_data, content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename={zip_filename}'
            return response
            
        except Exception as e:
            return Response(
                {'detail': f'Error during document generation: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
